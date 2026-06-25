"""
Capstone Report Auto-Filler
Project: AI Based School Bus Route Optimization

This script uses `python-docx` to load the template document, replace placeholders
for the project title, department, and course, and append full academic contents
for Chapters 1-6, including algorithms, pseudo-codes, and screenshots descriptions.
"""

import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def fill_report():
    doc_path = r"c:\School_Bus_CFAI\Y25_Capstone_Project_Report_Template.docx"
    output_path = r"c:\School_Bus_CFAI\Y25_Capstone_Project_Report_Template.docx"  # Overwrite original template

    if not os.path.exists(doc_path):
        print("Error: Template not found at", doc_path)
        return

    doc = docx.Document(doc_path)

    # =====================================================================
    # Part 1: Replace Placeholders in existing text
    # =====================================================================
    
    project_title = "AI-BASED SCHOOL BUS ROUTE OPTIMIZATION USING INTELLIGENT SEARCH AND PROBABILISTIC REASONING"
    department_name = "Department of Computer Science and Engineering"
    course_info = "25SC136E, Computational Foundations For AI"
    
    # We will search and replace in paragraphs
    for para in doc.paragraphs:
        # Title replacement
        if "<Title of the Project in Times New Roman, Font size: 18>" in para.text:
            para.text = para.text.replace("<Title of the Project in Times New Roman, Font size: 18>", project_title)
            para.style.font.name = 'Times New Roman'
            para.style.font.size = Pt(18)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        if " <Title of the Project> " in para.text:
            para.text = para.text.replace(" <Title of the Project> ", f"“ {project_title} ”")
        if " <Title of the Report > " in para.text:
            para.text = para.text.replace(" <Title of the Report > ", f"“ {project_title} ”")
        if "<Title of the Project>" in para.text:
            para.text = para.text.replace("<Title of the Project>", project_title)
        if "<Title of the Report >" in para.text:
            para.text = para.text.replace("<Title of the Report >", project_title)
            
        # Student List placeholder
        if "< List of the Students >" in para.text or "<List of Students >" in para.text:
            # We don't fill names, just cleaner placeholders
            para.text = para.text.replace("< List of the Students >", "[Student Name 1] ([Roll No 1])\n[Student Name 2] ([Roll No 2])\n[Student Name 3] ([Roll No 3])\n[Student Name 4] ([Roll No 4])")
            para.text = para.text.replace("<List of Students >", "[Student Name 1], [Student Name 2], [Student Name 3], [Student Name 4]")
            
        # Supervisor
        if "<Names and Designation of the Supervisor >" in para.text:
            para.text = para.text.replace("<Names and Designation of the Supervisor >", "[Supervisor Name]\n[Supervisor Designation]")
            
        # Department
        if "Department of <Name of the Discipline>" in para.text:
            para.text = para.text.replace("Department of <Name of the Discipline>", department_name)
            
    print("Placeholder replacements completed.")

    # =====================================================================
    # Part 2: Append Chapter Content to the end of the document
    # =====================================================================
    
    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        return p

    def add_body_text(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        return p

    def add_code_block(text):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(text)
        run.font.size = Pt(9.5)
        run.font.name = 'Courier New'
        return p

    # Insert a page break before starting chapter content
    doc.add_page_break()

    # --- CHAPTER 1 ---
    add_heading_1("CHAPTER 1: INTRODUCTION")
    add_body_text(
        "Efficient and safe transport planning for students is a major challenge for school administrations. "
        "Manual planning fails to handle variables like bus capacity limits, siblings requiring the same bus, "
        "and traffic congestion. This capstone project develops an AI-Based School Bus Route Optimization system. "
        "The system partitions students among buses under constraints, calculates optimal routes, and adapts to "
        "uncertain traffic conditions."
    )
    add_body_text(
        "The core objectives of this system are:\n"
        "1. Minimize travel time and distance for school buses.\n"
        "2. Ensure student safety and respect constraints (e.g. bus capacity, wheelchair accessibility).\n"
        "3. Incorporate uncertainty modeling to adapt to dynamic traffic events (weather, accidents).\n"
        "4. Provide explainable reasoning traces for all routing and scheduling decisions."
    )

    # --- CHAPTER 2 ---
    add_heading_1("CHAPTER 2: SOFTWARE DETAILS")
    add_body_text(
        "The system is built using a modern, decoupled web architecture consisting of a Python backend and a React frontend."
    )
    add_heading_2("Backend Technologies:")
    add_body_text(
        "• Python 3.13: Used for implementing the core artificial intelligence and search algorithms.\n"
        "• FastAPI: A modern, high-performance web framework for building REST APIs in Python. Exposes the algorithms to the client.\n"
        "• Uvicorn: An ASGI web server implementation for running FastAPI.\n"
        "• Python-Docx: Used for automated report generation."
    )
    
    add_heading_2("Frontend Technologies:")
    add_body_text(
        "• Vite React: Build tool and frontend library for constructing a fast, interactive user dashboard.\n"
        "• Recharts: A composable charting library for rendering empirical algorithm profiling graphs (runtime, memory, expansions).\n"
        "• Lucide-React: Icon library for the user interface.\n"
        "• Vanilla CSS: Configured for glassmorphic styling, dark-mode layouts, and map animations."
    )

    # --- CHAPTER 3 ---
    add_heading_1("CHAPTER 3: DESIGN & IMPLEMENTATION")
    
    add_heading_2("CO1: Agent Model & Problem Formulation")
    add_body_text(
        "The school bus routing problem is formulated as a single-agent system. The PEAS (Performance, Environment, Actuators, Sensors) "
        "specification is defined as follows:\n"
        "• Performance Measure: Minimize travel time, minimize distance, ensure student safety (respect capacity), and maximize punctuality.\n"
        "• Environment: Road network (graph of stops), students (pickup requests), and traffic conditions (stochastic).\n"
        "• Actuators: Navigation controls (routing choices), door controls, and passenger loading operations.\n"
        "• Sensors: GPS tracking, odometer, passenger counters, and traffic feeds.\n\n"
        "Problem Formulation: A state is defined by the current bus location, list of boarded students, visited nodes, and elapsed time. "
        "Actions represent moving to adjacent stops, picking up students, or dropping students at the school. The transition model "
        "predicts S' given (S, A). The path cost accumulates the travel times of all edges."
    )

    add_heading_2("CO2: Classical Search Algorithms")
    add_body_text(
        "To find the shortest paths between stop locations, several classical search algorithms are implemented: "
        "Breadth-First Search (BFS), Depth-First Search (DFS), Uniform Cost Search (UCS), A* Search, Greedy Best-First Search, "
        "and a memory-bounded variant (IDA*). The heuristic function h(n) computes the scaled straight-line Euclidean distance "
        "from stop coordinates to the destination. Since straight-line distance is the shortest possible path, h(n) is "
        "guaranteed to be admissible (never overestimates cost) and consistent (satisfies the triangle inequality)."
    )

    add_heading_2("CO3: Constraint Satisfaction Problems (CSP)")
    add_body_text(
        "Student allocation to buses is modeled as a CSP. The variables represent students, and domains represent the "
        "available buses (Bus 1, Bus 2). The constraints are: (1) capacity constraints (max 8 students per bus), "
        "(2) wheelchair constraints (wheelchair-reliant student Chloe must ride Bus 1), and (3) sibling constraints "
        "(siblings must be assigned to the same bus). The CSP is solved using Backtracking Search with MRV (Minimum "
        "Remaining Values), Degree Heuristics, LCV (Least Constraining Value), and Forward Checking constraint propagation. "
        "We also implement a Min-Conflicts local search solver and show how the problem maps to SAT CNF clauses."
    )

    add_heading_2("CO4: Adversarial & Stochastic Decisions")
    add_body_text(
        "We evaluate routing under worst-case and expected traffic conditions. Minimax Search with Alpha-Beta Pruning is used "
        "to compute robust paths against a worst-case 'Traffic Congestion' adversary that injects travel delays. Bounded rationality "
        "is implemented via a depth limit and an evaluation function estimating remaining travel time. Expectimax search is "
        "implemented to model stochastic traffic conditions, averaging travel times over probability distributions."
    )

    add_heading_2("CO5: Probabilistic Reasoning under Uncertainty")
    add_body_text(
        "A Bayesian Network models travel delays. Variables include Weather (W), Accident (A), RoadWork (RW), TrafficCongestion (C), "
        "and BusDelay (D). Exact inference is implemented using Variable Elimination. Approximate inference is modeled using Likelihood "
        "Weighting. A Hidden Markov Model (HMM) tracks the true congestion level of a road segment over time from noisy GPS "
        "velocity observations using Forward Filtering. Expected utility decisions are evaluated to select optimal routes."
    )

    add_heading_2("CO6: Hybrid Architecture")
    add_body_text(
        "The master system integrates all components into a hybrid architecture. First, the CSP engine assigns students to buses. "
        "Next, A* search generates optimal paths. The Bayesian Network evaluates traffic conditions and predicts delays. The decision logic "
        "reroutes the buses if delay probabilities are high. A detailed failure analysis evaluates recovery strategies for bus breakdowns, "
        "verifying capacity limits. Ethics checks address heuristic bias (where outer regions are systematically penalized with longer travel times) "
        "and calibration errors."
    )

    # --- ALGORITHMS & PSEUDO CODES ---
    add_heading_1("ALGORITHMS & PSEUDO CODES")
    
    add_heading_2("1. A* Search Algorithm")
    add_code_block(
        "Algorithm A*(start, goal):\n"
        "  Initialize OpenSet (PriorityQueue) with (start, f_score = h(start))\n"
        "  Initialize g_score[start] = 0\n"
        "  Initialize came_from (map to reconstruct path)\n\n"
        "  while OpenSet is not empty:\n"
        "    current = OpenSet.pop()\n"
        "    if current == goal:\n"
        "      return reconstruct_path(came_from, current)\n\n"
        "    for each neighbor of current with edge_weight c:\n"
        "      tentative_g = g_score[current] + c\n"
        "      if tentative_g < g_score[neighbor]:\n"
        "        came_from[neighbor] = current\n"
        "        g_score[neighbor] = tentative_g\n"
        "        f_score = tentative_g + h(neighbor)\n"
        "        OpenSet.push(neighbor, f_score)\n\n"
        "  return Failure"
    )

    add_heading_2("2. CSP Backtracking Search")
    add_code_block(
        "Algorithm Backtrack(assignment, csp):\n"
        "  if assignment is complete:\n"
        "    return assignment\n\n"
        "  var = Select-Unassigned-Variable(csp, assignment)  # Uses MRV / Degree Heuristics\n"
        "  for each value in Order-Domain-Values(var, assignment, csp):  # Uses LCV Heuristic\n"
        "    if Value is consistent with assignment according to Constraints:\n"
        "      Add {var = value} to assignment\n"
        "      pruned_domains = ForwardCheck(var, value, assignment, csp)  # Constraint propagation\n"
        "      if pruned_domains is not None:\n"
        "        result = Backtrack(assignment, csp)\n"
        "        if result is not None:\n"
        "          return result\n"
        "      Remove {var = value} from assignment and restore pruned_domains\n"
        "  return Failure"
    )

    add_heading_2("3. Bayesian Variable Elimination")
    add_code_block(
        "Algorithm VariableElimination(QueryVar, Evidence, HiddenVars, Factors):\n"
        "  # Step 1: Restrict factors based on evidence\n"
        "  For each factor F in Factors:\n"
        "    Restrict variables in F matching Evidence\n\n"
        "  # Step 2: Eliminate hidden variables\n"
        "  For each variable V in HiddenVars:\n"
        "    DependentFactors = factors containing V\n"
        "    IndependentFactors = factors not containing V\n"
        "    ProductFactor = Multiply(DependentFactors)\n"
        "    SummedFactor = SumOut(V, ProductFactor)\n"
        "    Factors = IndependentFactors + [SummedFactor]\n\n"
        "  # Step 3: Multiply remaining factors and normalize\n"
        "  FinalFactor = Multiply(Factors)\n"
        "  return Normalize(FinalFactor)"
    )

    # --- CHAPTER 4 ---
    add_heading_1("CHAPTER 4: OUTPUT SCREENS")
    add_body_text(
        "The frontend Vite React user interface provides a dashboard to monitor and interact with the AI algorithms:\n"
        "1. Dashboard Overview: Displays a visual map of the school bus network showing nodes and student locations. "
        "KPI cards show total students, active buses, weather conditions, and average traffic delay.\n"
        "2. CO1 Agent Model Tab: Simulates the step-by-step state transitions of a school bus. Displays state variables "
        "(current node, passenger list, elapsed time) and outputs an explainable trace log.\n"
        "3. CO2 Classical Search Tab: Animates the search frontier expansions of BFS, DFS, UCS, A*, and Greedy. "
        "Includes empirical profiling charts comparing runtime, node expansions, and peak memory usage.\n"
        "4. CO3 CSP Scheduling Tab: Visualizes student-to-bus allocations. Shows backtracking steps and includes "
        "an explainability panel explaining exactly why a constraint failed.\n"
        "5. CO4 Adversarial Search Tab: Compares worst-case Minimax Alpha-Beta routing against stochastic Expectimax routing, "
        "illustrating policies under bounded rationality.\n"
        "6. CO5 Probabilistic Tab: Provides controls to set weather and accident evidence, dynamically updating Bayesian network "
        "inference. Tracks bus state via HMM filtering graphs.\n"
        "7. CO6 Hybrid Tab: Visualizes the combined morning operation. Includes failure recovery controls (triggering a bus breakdown) "
        "and lists the ethical analysis."
    )

    # --- CHAPTER 5 ---
    add_heading_1("CHAPTER 5: CONCLUSION")
    add_body_text(
        "This project demonstrates the application of key computational foundations of artificial intelligence to school bus routing. "
        "By integrating CSP constraint solvers, search pathfinders, Bayesian Networks, HMM state tracking, and decision engines, the system "
        "achieves efficient scheduling and routing that dynamically adapts to traffic uncertainty. While classical search and heuristics "
        "optimize paths, technical limitations such as heuristic bias and CPT calibration error must be carefully calibrated to ensure "
        "fairness and student safety. Future developments include integrating real-world geographic coordinates (GIS) and optimizing electric "
        "bus battery constraints."
    )

    # --- CHAPTER 6 ---
    add_heading_1("CHAPTER 6: REFERENCE")
    add_body_text(
        "1. Russell, S., & Norvig, P. (2020). Artificial Intelligence: A Modern Approach (4th ed.). Pearson.\n"
        "2. Cormen, T. H., Leiserson, C. E., Rivest, R. L., & Stein, C. (2009). Introduction to Algorithms (3rd ed.). MIT Press.\n"
        "3. Koller, D., & Friedman, N. (2009). Probabilistic Graphical Models: Principles and Techniques. MIT Press.\n"
        "4. Ghallab, M., Nau, D., & Traverso, P. (2004). Automated Planning: Theory and Practice. Morgan Kaufmann."
    )

    doc.save(output_path)
    print("Report written successfully to", output_path)

if __name__ == "__main__":
    fill_report()
